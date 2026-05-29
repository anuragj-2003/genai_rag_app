import os
import tempfile
import pandas as pd
import pypdf
import docx
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter

def process_uploaded_file(uploaded_file):
    """
    Processes an uploaded file (PDF, DOCX, XLSX) and returns a list of LlamaIndex Documents.
    
    Input:
        uploaded_file (UploadedFile): An object containing .name and .getvalue()
        
    Output:
        list: A list of LlamaIndex Document objects with metadata (source, page).
    """
    if uploaded_file is None:
        return []

    file_extension = os.path.splitext(uploaded_file.name)[1].lower()
    documents = []

    # Create a temporary file to save the uploaded content
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    try:
        if file_extension == ".pdf":
            reader = pypdf.PdfReader(tmp_file_path)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    documents.append(Document(
                        text=page_text,
                        metadata={"source": uploaded_file.name, "page": i + 1}
                    ))
        elif file_extension == ".docx":
            doc_file = docx.Document(tmp_file_path)
            full_text = []
            for para in doc_file.paragraphs:
                full_text.append(para.text)
            text = '\n'.join(full_text)
            if text.strip():
                documents.append(Document(
                    text=text,
                    metadata={"source": uploaded_file.name, "page": 1}
                ))
        elif file_extension in [".xlsx", ".xls"]:
            try:
                df = pd.read_excel(tmp_file_path)
                text_content = df.to_string()
                if text_content.strip():
                    documents.append(Document(
                        text=text_content,
                        metadata={"source": uploaded_file.name, "page": 1}
                    ))
            except Exception:
                pass
        else:
            # Fallback for text files
            try:
                with open(tmp_file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                if text.strip():
                    documents.append(Document(
                        text=text,
                        metadata={"source": uploaded_file.name, "page": 1}
                    ))
            except Exception:
                pass
    finally:
        # Clean up temp file
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)

    # Split text if we have documents
    if documents:
        # Use LlamaIndex SentenceSplitter
        splitter = SentenceSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        nodes = splitter.get_nodes_from_documents(documents)
        
        # Convert nodes to LlamaIndex Document objects for indexing compatibility
        split_docs = []
        for node in nodes:
            # Preserve/set metadata
            source = node.metadata.get("source", uploaded_file.name)
            page = node.metadata.get("page", 1)
            
            split_docs.append(Document(
                text=node.text,
                metadata={
                    "source": source,
                    "page": page
                }
            ))
        return split_docs
    
    return []
